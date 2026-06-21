"""多格式内容解析器，支持 PDF/Word/HTML/TXT/CSV/JSON 解析与长文档拆分.

采用可选依赖策略：未安装对应库时优雅降级到基础文本提取。
"""

import csv
import io
import json
import re
from pathlib import Path
from typing import List, Optional, Tuple

# 可选依赖检测
try:
    import pypdf  # type: ignore
    PDF_AVAILABLE = True
except ImportError:
    try:
        import PyPDF2  # type: ignore  # noqa: N813
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False

try:
    import docx  # type: ignore  # python-docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup  # type: ignore
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False


class MultiFormatParser:
    """多格式文档解析器，统一输出 Markdown 文本."""

    SUPPORTED_EXTENSIONS = {
        '.md', '.markdown', '.txt', '.csv', '.json', '.html', '.htm',
        '.pdf', '.docx', '.doc', '.rst', '.log'
    }

    def __init__(self):
        self.warnings: List[str] = []

    def parse(self, file_path: Path) -> Tuple[str, str]:
        """解析文件，返回 (title, markdown_content).

        Args:
            file_path: 文件路径

        Returns:
            (标题, Markdown 格式内容) 元组
        """
        suffix = file_path.suffix.lower()
        title = file_path.stem.replace('_', ' ').replace('-', ' ').title()

        if suffix in ('.md', '.markdown'):
            content = file_path.read_text(encoding='utf-8', errors='ignore')
            extracted_title = self._extract_md_title(content)
            if extracted_title:
                title = extracted_title
            return title, content

        if suffix == '.txt':
            content = file_path.read_text(encoding='utf-8', errors='ignore')
            return title, f"# {title}\n\n{content}"

        if suffix == '.rst':
            content = file_path.read_text(encoding='utf-8', errors='ignore')
            return title, f"# {title}\n\n{self._rst_to_md(content)}"

        if suffix == '.log':
            content = file_path.read_text(encoding='utf-8', errors='ignore')
            return title, f"# {title}\n\n```\n{content}\n```"

        if suffix == '.json':
            return self._parse_json(file_path, title)

        if suffix == '.csv':
            return self._parse_csv(file_path, title)

        if suffix in ('.html', '.htm'):
            return self._parse_html(file_path, title)

        if suffix == '.pdf':
            return self._parse_pdf(file_path, title)

        if suffix == '.docx':
            return self._parse_docx(file_path, title)

        if suffix == '.doc':
            self.warnings.append(f".doc 旧格式不支持（需转换到 .docx）: {file_path}")
            return title, f"# {title}\n\n> 警告：.doc 旧格式不支持，请转换为 .docx"

        # 未知格式，尝试按文本读取
        try:
            content = file_path.read_text(encoding='utf-8', errors='ignore')
            return title, f"# {title}\n\n{content}"
        except (IOError, OSError) as e:
            self.warnings.append(f"无法读取文件 {file_path}: {e}")
            return title, f"# {title}\n\n> 读取失败: {e}"

    def _parse_json(self, file_path: Path, title: str) -> Tuple[str, str]:
        """解析 JSON 文件，转换为 Markdown 表格或列表."""
        try:
            data = json.loads(file_path.read_text(encoding='utf-8'))
        except (json.JSONDecodeError, UnicodeDecodeError) as e:
            return title, f"# {title}\n\n> JSON 解析失败: {e}"

        parts = [f"# {title}\n"]
        if isinstance(data, list):
            parts.append(f"共 {len(data)} 条记录\n")
            if data and isinstance(data[0], dict):
                # 转表格
                keys = list(data[0].keys())
                parts.append('| ' + ' | '.join(keys) + ' |')
                parts.append('| ' + ' | '.join(['---'] * len(keys)) + ' |')
                for item in data[:100]:  # 限制 100 行
                    parts.append('| ' + ' | '.join(str(item.get(k, '')) for k in keys) + ' |')
            else:
                for item in data[:100]:
                    parts.append(f"- {item}")
        elif isinstance(data, dict):
            parts.append("\n| 字段 | 值 |")
            parts.append("| --- | --- |")
            for k, v in data.items():
                val_str = json.dumps(v, ensure_ascii=False) if isinstance(v, (dict, list)) else str(v)
                if len(val_str) > 200:
                    val_str = val_str[:200] + '...'
                parts.append(f"| {k} | {val_str} |")
        else:
            parts.append(f"\n```\n{json.dumps(data, indent=2, ensure_ascii=False)}\n```")

        return title, '\n'.join(parts)

    def _parse_csv(self, file_path: Path, title: str) -> Tuple[str, str]:
        """解析 CSV 文件，转换为 Markdown 表格."""
        try:
            content = file_path.read_text(encoding='utf-8', errors='ignore')
            reader = csv.reader(io.StringIO(content))
            rows = list(reader)
        except (csv.Error, UnicodeDecodeError) as e:
            return title, f"# {title}\n\n> CSV 解析失败: {e}"

        if not rows:
            return title, f"# {title}\n\n> 空文件"

        parts = [f"# {title}\n"]
        parts.append(f"共 {len(rows)} 行\n")
        # 表头
        header = rows[0]
        parts.append('| ' + ' | '.join(header) + ' |')
        parts.append('| ' + ' | '.join(['---'] * len(header)) + ' |')
        # 数据行（限制 100 行）
        for row in rows[1:101]:
            # 补齐列数
            while len(row) < len(header):
                row.append('')
            parts.append('| ' + ' | '.join(row[:len(header)]) + ' |')
        if len(rows) > 101:
            parts.append(f"\n> ... 共 {len(rows) - 1} 行数据，仅显示前 100 行")

        return title, '\n'.join(parts)

    def _parse_html(self, file_path: Path, title: str) -> Tuple[str, str]:
        """解析 HTML 文件，提取文本内容."""
        html_content = file_path.read_text(encoding='utf-8', errors='ignore')

        if BS4_AVAILABLE:
            soup = BeautifulSoup(html_content, 'html.parser')
            # 提取标题
            if soup.title and soup.title.string:
                title = soup.title.string.strip()
            # 移除脚本和样式
            for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
                tag.decompose()
            # 转换为 Markdown 风格
            parts = [f"# {title}\n"]
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'pre', 'code', 'blockquote']):
                text = element.get_text(strip=True)
                if not text:
                    continue
                tag_name = element.name
                if tag_name == 'h1':
                    parts.append(f"\n## {text}\n")
                elif tag_name == 'h2':
                    parts.append(f"\n### {text}\n")
                elif tag_name in ('h3', 'h4', 'h5', 'h6'):
                    parts.append(f"\n#### {text}\n")
                elif tag_name == 'li':
                    parts.append(f"- {text}")
                elif tag_name == 'pre':
                    parts.append(f"\n```\n{text}\n```\n")
                elif tag_name == 'blockquote':
                    parts.append(f"\n> {text}\n")
                else:
                    parts.append(text + '\n')
            return title, '\n'.join(parts)
        else:
            # 无 BeautifulSoup，简单正则提取
            # 提取 title
            title_match = re.search(r'<title[^>]*>(.*?)</title>', html_content, re.IGNORECASE | re.DOTALL)
            if title_match:
                title = title_match.group(1).strip()
            # 移除标签
            text = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.IGNORECASE | re.DOTALL)
            text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.IGNORECASE | re.DOTALL)
            text = re.sub(r'<[^>]+>', ' ', text)
            text = re.sub(r'\s+', ' ', text).strip()
            return title, f"# {title}\n\n{text}"

    def _parse_pdf(self, file_path: Path, title: str) -> Tuple[str, str]:
        """解析 PDF 文件，提取文本."""
        if not PDF_AVAILABLE:
            self.warnings.append(
                "PDF 解析需要 pypdf: pip install pypdf"
            )
            return title, f"# {title}\n\n> PDF 解析需要安装 pypdf: `pip install pypdf`"

        try:
            parts = [f"# {title}\n"]
            if 'pypdf' in globals():
                reader = pypdf.PdfReader(str(file_path))
            else:
                reader = PyPDF2.PdfReader(str(file_path))  # noqa: F821

            num_pages = len(reader.pages)
            parts.append(f"共 {num_pages} 页\n")

            # 尝试从元数据提取标题
            if reader.metadata:
                meta_title = reader.metadata.get('/Title', '')
                if meta_title:
                    title = str(meta_title)
                    parts[0] = f"# {title}\n"

            for i, page in enumerate(reader.pages, 1):
                page_text = page.extract_text() or ''
                if page_text.strip():
                    parts.append(f"\n## 第 {i} 页\n\n{page_text.strip()}\n")

            return title, '\n'.join(parts)
        except Exception as e:  # noqa: BLE001
            self.warnings.append(f"PDF 解析失败: {e}")
            return title, f"# {title}\n\n> PDF 解析失败: {e}"

    def _parse_docx(self, file_path: Path, title: str) -> Tuple[str, str]:
        """解析 Word .docx 文件，提取文本."""
        if not DOCX_AVAILABLE:
            self.warnings.append(
                "Word 解析需要 python-docx: pip install python-docx"
            )
            return title, f"# {title}\n\n> Word 解析需要安装 python-docx: `pip install python-docx`"

        try:
            doc = docx.Document(str(file_path))  # noqa: F821
            parts = [f"# {title}\n"]

            # 从核心属性提取标题
            if doc.core_properties.title:
                title = doc.core_properties.title
                parts[0] = f"# {title}\n"

            para_count = 0
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style = para.style.name.lower() if para.style else ''
                if 'heading 1' in style:
                    parts.append(f"\n## {text}\n")
                elif 'heading 2' in style:
                    parts.append(f"\n### {text}\n")
                elif 'heading' in style:
                    parts.append(f"\n#### {text}\n")
                else:
                    parts.append(text + '\n')
                para_count += 1

            # 提取表格
            for table in doc.tables:
                parts.append('\n')
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    parts.append('| ' + ' | '.join(cells) + ' |')

            return title, '\n'.join(parts)
        except Exception as e:  # noqa: BLE001
            self.warnings.append(f"Word 解析失败: {e}")
            return title, f"# {title}\n\n> Word 解析失败: {e}"

    def _extract_md_title(self, content: str) -> Optional[str]:
        """从 Markdown 内容提取标题."""
        match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        if match:
            return match.group(1).strip()
        if content.startswith('---'):
            parts = content.split('---', 2)
            if len(parts) >= 3:
                fm_match = re.search(r'^title:\s*(.+)$', parts[1], re.MULTILINE)
                if fm_match:
                    return fm_match.group(1).strip().strip('\'"')
        return None

    def _rst_to_md(self, content: str) -> str:
        """简单 RST 转 Markdown."""
        lines = content.split('\n')
        result = []
        for line in lines:
            # RST 标题下划线转换
            if line.strip() and all(c == '=' for c in line.strip()):
                if result:
                    result[-1] = f"## {result[-1]}"
                continue
            if line.strip() and all(c == '-' for c in line.strip()):
                if result:
                    result[-1] = f"### {result[-1]}"
                continue
            result.append(line)
        return '\n'.join(result)


class LongDocumentSplitter:
    """长文档自动拆分为多个知识原子.

    按章节标题或字数阈值拆分。
    """

    def __init__(self, max_chars_per_atom: int = 2000, min_chars_per_atom: int = 200):
        self.max_chars = max_chars_per_atom
        self.min_chars = min_chars_per_atom

    def split(self, title: str, content: str) -> List[Tuple[str, str]]:
        """拆分长文档为多个 (子标题, 子内容) 元组.

        Args:
            title: 文档主标题
            content: 文档 Markdown 内容

        Returns:
            (子标题, 子内容) 列表
        """
        # 先按章节拆分
        sections = self._split_by_headings(content)

        # 对超长章节再按字数拆分
        result = []
        for section_title, section_content in sections:
            if len(section_content) <= self.max_chars:
                result.append((section_title, section_content))
            else:
                # 按段落拆分
                chunks = self._split_by_paragraphs(section_content)
                for i, chunk in enumerate(chunks, 1):
                    chunk_title = f"{section_title}（{i}）" if len(chunks) > 1 else section_title
                    result.append((chunk_title, chunk))

        # 合并过短的片段
        result = self._merge_short_sections(result)

        return result if result else [(title, content)]

    def _split_by_headings(self, content: str) -> List[Tuple[str, str]]:
        """按 Markdown 标题拆分."""
        sections = []
        current_title = ''
        current_lines: List[str] = []

        for line in content.split('\n'):
            if re.match(r'^#{1,3}\s+', line):
                # 保存前一节
                if current_lines:
                    section_content = '\n'.join(current_lines).strip()
                    if section_content:
                        sections.append((current_title, section_content))
                current_title = re.sub(r'^#+\s+', '', line).strip()
                current_lines = [line]
            else:
                current_lines.append(line)

        # 最后一节
        if current_lines:
            section_content = '\n'.join(current_lines).strip()
            if section_content:
                sections.append((current_title, section_content))

        return sections

    def _split_by_paragraphs(self, content: str) -> List[str]:
        """按段落拆分超长内容."""
        paragraphs = content.split('\n\n')
        chunks = []
        current_chunk = ''
        for para in paragraphs:
            if len(current_chunk) + len(para) > self.max_chars and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = para
            else:
                current_chunk = current_chunk + '\n\n' + para if current_chunk else para
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        return chunks

    def _merge_short_sections(self, sections: List[Tuple[str, str]]) -> List[Tuple[str, str]]:
        """合并过短的章节."""
        if not sections:
            return sections
        merged = []
        buffer_title = ''
        buffer_content = ''

        for title, content in sections:
            if len(content) < self.min_chars:
                if buffer_content:
                    buffer_content += '\n\n' + content
                else:
                    buffer_title = title
                    buffer_content = content
            else:
                if buffer_content:
                    merged.append((buffer_title or title, buffer_content))
                    buffer_title = ''
                    buffer_content = ''
                merged.append((title, content))

        if buffer_content:
            if merged:
                last_title, last_content = merged[-1]
                merged[-1] = (last_title, last_content + '\n\n' + buffer_content)
            else:
                merged.append((buffer_title, buffer_content))

        return merged


def get_supported_formats_info() -> str:
    """返回支持的格式信息字符串."""
    info = [
        "支持格式：",
        f"  ✅ Markdown (.md/.markdown) - 原生支持",
        f"  ✅ 纯文本 (.txt/.rst/.log) - 原生支持",
        f"  ✅ JSON (.json) - 原生支持",
        f"  ✅ CSV (.csv) - 原生支持",
        f"  {'✅' if BS4_AVAILABLE else '⚠️'} HTML (.html/.htm) - {'已安装' if BS4_AVAILABLE else '需 beautifulsoup4'}",
        f"  {'✅' if PDF_AVAILABLE else '⚠️'} PDF (.pdf) - {'已安装' if PDF_AVAILABLE else '需 pypdf'}",
        f"  {'✅' if DOCX_AVAILABLE else '⚠️'} Word (.docx) - {'已安装' if DOCX_AVAILABLE else '需 python-docx'}",
    ]
    return '\n'.join(info)
